import os
import random

# package python-docx
from docx import Document
from docx.shared import Pt  # For setting font size
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# get word_dir
work_dir = os.path.dirname(os.path.abspath(__file__))
# print(work_dir)

# define doc variables
source_doc = Document(work_dir + '/source.docx')
target_doc = Document(work_dir + '/template.docx')

# style control
font_name = [u'对你不止是喜欢', u'你是我的可爱宝贝', u'品如手写体']
font_size = [16, 15, 16, 17, 16]

# pasting words into target_docx
for para in source_doc.paragraphs:
    # iterating every paragraph
    new_para = target_doc.add_paragraph()
    for char in para.text:
        # iterating every character
        run = new_para.add_run(char)
        # 1. set font name(type)
        font_type = random.randint(0, 2)
        run.font.name = font_name[font_type]
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name[font_type])
        # 2. set font size
        run.font.size = Pt(font_size[random.randint(0, 4)])
        # 3. set font italic
        # run.font.italic = True if random.randint(0, 5) == 1 else False
        # 4. set font bold
        # run.font.bold = True if random.randint(0, 5) == 1 else False
        # 5. set font spacing and vertical offset
        spacing = OxmlElement('w:spacing')
        vertical = OxmlElement('w:position')
        spacing.set(qn('w:val'), str(random.randint(1, 15) - 7))
        vertical.set(qn('w:val'), str(random.randint(1, 5) - 2))  # Adjust value for vertical offset
        # Add the spacing element to the run properties
        if run._element.rPr is None:
            run._element.rPr = OxmlElement('w:rPr')
        run._element.rPr.append(spacing)
        run._element.rPr.append(vertical)
# 保存目标文档
target_doc.save(work_dir + '/template.docx')
